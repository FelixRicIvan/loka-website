# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()

    # ---------------------------------------------------------------------------
    # 版面設定 (Margins)
    # ---------------------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---------------------------------------------------------------------------
    # 品牌視覺色系 (LOKA Brand Colors)
    # ---------------------------------------------------------------------------
    COLOR_PRIMARY = RGBColor(10, 59, 46)     # 奢華深綠 (#0A3B2E)
    COLOR_SECONDARY = RGBColor(17, 77, 62)   # 亮深綠 (#114D3E)
    COLOR_ACCENT = RGBColor(226, 241, 237)   # 淺綠白 (#E2F1ED)
    COLOR_TEXT = RGBColor(51, 51, 51)        # 暗灰內文 (#333333)
    COLOR_MUTED = RGBColor(102, 102, 102)    # 灰色 (#666666)

    # ---------------------------------------------------------------------------
    # 輔育設定函數
    # ---------------------------------------------------------------------------
    def set_font(run, font_name="Segoe UI", font_name_east_asia="微軟正黑體", size_pt=11, bold=False, italic=False, color=COLOR_TEXT):
        run.font.name = font_name
        # 設定東亞中文字型
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name_east_asia)
        rPr.append(rFonts)
        
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, size_pt=16, bold=True, color=COLOR_PRIMARY)
        # 加上底部分割線
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                         r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="0A3B2E"/>'
                         r'</w:pBdr>')
        pPr.append(pBdr)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, size_pt=13, bold=True, color=COLOR_SECONDARY)
        return p

    def add_paragraph(text, bold_prefix="", indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if indent > 0:
            p.paragraph_format.left_indent = Inches(indent)
        
        if bold_prefix:
            run_pfx = p.add_run(bold_prefix)
            set_font(run_pfx, size_pt=10.5, bold=True)
            
        run = p.add_run(text)
        set_font(run, size_pt=10.5, color=COLOR_TEXT)
        return p

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shading)

    # ---------------------------------------------------------------------------
    # 封面設計 (Cover Page)
    # ---------------------------------------------------------------------------
    # 頂部大區塊空白
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(80)
    
    # LOKA 主標題
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("L O K A")
    set_font(run_logo, size_pt=36, bold=True, color=COLOR_PRIMARY)
    p_logo.paragraph_format.space_after = Pt(0)
    
    # 奢華副標
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Helping you uncover the beauty of the world.")
    set_font(run_sub, font_name="Georgia", font_name_east_asia="微軟正黑體", size_pt=10, italic=True, color=COLOR_MUTED)
    p_sub.paragraph_format.space_after = Pt(30)

    # 橫幅裝飾線
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("—" * 35)
    set_font(run_line, size_pt=10, color=COLOR_PRIMARY)
    p_line.paragraph_format.space_after = Pt(40)

    # 報告主標題
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("輕量化豪華旅行規劃助手\n專案結案與系統設計報告")
    set_font(run_title, size_pt=20, bold=True, color=COLOR_PRIMARY)
    p_title.paragraph_format.space_after = Pt(120)

    # 專案中繼資訊 (Metadata)
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.line_spacing = 1.3
    
    runs_meta = [
        ("專案名稱：", "LOKA Bespoke Luxury Travel Planner\n"),
        ("開發階段：", "系統佈局最佳化與代碼重構階段 (完成)\n"),
        ("報告日期：", "2026 年 6 月 11 日\n"),
        ("開發團隊：", "Antigravity Coding Assistant & Partner")
    ]
    for label, val in runs_meta:
        r_lbl = p_meta.add_run(label)
        set_font(r_lbl, size_pt=10, bold=True, color=COLOR_MUTED)
        r_val = p_meta.add_run(val)
        set_font(r_val, size_pt=10, color=COLOR_TEXT)
        
    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # 一、專案概述
    # ---------------------------------------------------------------------------
    add_heading_1("一、 專案概述 (Project Overview)")
    add_paragraph(
        "LOKA 是一款專為追求高質感行程的旅客量身打造的超輕量、奢華風格旅行規劃助理。透過整合 Google Gemini 大語言模型，"
        "LOKA 能夠根據旅客的目的地、出遊天數及旅伴人數，於秒級內自動產出結構完整、無景點重複循環的客製化旅行時間軸（Timeline）。"
        "此外，專案深度結合了 Leaflet 開源地圖與 OSRM (Open Source Routing Machine) 真實道路路徑規劃 API，"
        "使旅客不僅能以雜誌級的排版閱讀行程，還能在地圖上直觀查看真實的駕駛路線、預估行程花費，並自由將規劃儲存至雲端，或匯出為 "
        "Notion 筆記及高品質 PDF 手冊。"
    )
    add_paragraph(
        "本階段開發的重點在於：調整主介面的版面配置，滿足旅客對於「左側閱讀行程、右側對照地圖」的流暢操作期待；"
        "修復地圖因外層滾動容器干擾而失效的 Sticky 黏性定位 bug；新增橫縱軸排版切換功能（X/Y 軸對調），"
        "並對全站的 React 與 FastAPI 進行代碼重構與效能優化，全面消除無謂的重複渲染並健全依賴關係。"
    )

    # ---------------------------------------------------------------------------
    # 二、 系統架構設計
    # ---------------------------------------------------------------------------
    add_heading_1("二、 系統架構設計 (System Architecture)")
    add_paragraph("LOKA 採用現代前後端分離的極簡單頁應用 (SPA) 架構，其結構交互如下：")
    
    add_paragraph("使用者透過精緻的前端介面輸入旅行偏好，發送請求至 FastAPI 後端伺服器。", bold_prefix="1. 請求層 (Request Layer): ")
    add_paragraph("FastAPI 後端在驗證 JWT 身分憑證後，呼叫 AI 行程規劃生成引擎 (ai_engine.py)，將使用者的需求結合 System Instructions 送入 Google Gemini 大語言模型進行推理。", bold_prefix="2. 智慧生成層 (AI Layer): ")
    add_paragraph("AI 引擎透過 Pydantic Schema 強制規範輸出完全符合結構的 JSON 行程資料，以防格式錯誤導致崩潰。", bold_prefix="3. 資料流控層 (Data Layer): ")
    add_paragraph("前端接收資料後，由 App.jsx 渲染打卡時間軸，同時將地理座標陣列傳遞至 MapComponent.jsx，在地圖上動態標註景點，並呼叫外網 OSRM API 繪製高光道路軌跡線。", bold_prefix="4. 視覺化層 (Map Layer): ")

    add_heading_2("1. 專案目錄結構 (Directory Layout)")
    add_paragraph("專案的程式碼結構清晰、分工明確，主要檔案清單與功能如下：")
    
    catalog_table = doc.add_table(rows=7, cols=3)
    catalog_table.style = 'Light Shading Accent 1'
    hdr_cells = catalog_table.rows[0].cells
    hdr_cells[0].text = '檔案 / 目錄名稱'
    hdr_cells[1].text = '檔案類型'
    hdr_cells[2].text = '主要職責與功能說明'
    
    set_cell_background(hdr_cells[0], "0A3B2E")
    set_cell_background(hdr_cells[1], "0A3B2E")
    set_cell_background(hdr_cells[2], "0A3B2E")
    for r in range(3):
        set_font(hdr_cells[r].paragraphs[0].runs[0], size_pt=9.5, bold=True, color=RGBColor(255, 255, 255))
        
    table_data = [
        ("main.py", "Python 腳本", "FastAPI 後端伺服器，負責伺服前端 HTML/JSX 檔案、處理 JWT 登入註冊、儲存/刪除雲端行程與模擬付費金流網關。"),
        ("ai_engine.py", "Python 腳本", "LOKA 行程生成大腦。整合官方 Google Gen AI 客戶端，定義 Pydantic 結構架構（ItinerarySchema），呼叫 gemini-2.5-flash 推理並輸出結構化 JSON。"),
        ("index.html", "HTML 頁面", "前端主入口。載入 Google Outfit 字型、Leaflet 地圖庫、React 框架與 Babel Standalone 即時解析器，並設有全域 JS 執行期診斷健檢機制。"),
        ("App.jsx", "React / JSX", "前端核心組件。實現主畫面佈局、狀態管理（多語系、行程、付費牆、打卡狀態）、模擬支付卡片處理，並整合 PDF 與 Notion 行程匯出工具。"),
        ("MapComponent.jsx", "React / JSX", "地圖專屬組件。初始化 Leaflet 地圖、動態繪製 HTML+Tailwind 景點 Marker、監聽 focusedCoords 進行平滑飛越（flyTo），並串接 OSRM 真實路線 API。"),
        ("static/", "目錄", "本地靜態資源備份。存放 Leaflet CSS/JS 圖磚依賴、React 核心庫與 Tailwind JS，防止因無外網或 CDN 逾時導致專案崩潰。")
    ]
    
    for idx, (name, ftype, desc) in enumerate(table_data):
        row_cells = catalog_table.rows[idx+1].cells
        row_cells[0].text = name
        row_cells[1].text = ftype
        row_cells[2].text = desc
        
        # 美化表格文字字型
        for c_idx in range(3):
            set_font(row_cells[c_idx].paragraphs[0].runs[0], size_pt=9.5)
            if idx % 2 == 1:
                set_cell_background(row_cells[c_idx], "F8FAF9")

    # ---------------------------------------------------------------------------
    # 三、 資料庫設計
    # ---------------------------------------------------------------------------
    add_heading_2("2. 資料庫設計 (Database Schema)")
    add_paragraph("LOKA 後端使用 SQLite 作為雲端行程與會員資料持久化儲存媒介。資料庫中包含兩大實體資料表（SQLAlchemy 模型）：")
    
    add_paragraph("儲存註冊使用者的帳密與付費狀態。包含欄位：", bold_prefix="• User 表 (users 表): ")
    add_paragraph("主鍵，自動遞增。", bold_prefix="  - id (Integer, PK): ", indent=0.4)
    add_paragraph("唯一識別碼，用於登入與身分確認。", bold_prefix="  - username (String, Unique): ", indent=0.4)
    add_paragraph("使用者電子信箱（用於新版註冊）。", bold_prefix="  - email (String, Nullable): ", indent=0.4)
    add_paragraph("以 bcrypt 雜湊加密儲存的安全密碼。", bold_prefix="  - password_hash (String): ", indent=0.4)
    add_paragraph("布林值旗標，標記使用者是否解鎖了 Premium 尊榮會員權限（匯出與雲端功能）。", bold_prefix="  - is_premium (Boolean, Default=False): ", indent=0.4)

    add_paragraph("儲存使用者儲存的客製化豪華旅行時間軸。包含欄位：", bold_prefix="• ItineraryModel 表 (itineraries 表): ")
    add_paragraph("主鍵，自動遞增。", bold_prefix="  - id (Integer, PK): ", indent=0.4)
    add_paragraph("外鍵，關聯至 users.id，用於區分各個使用者的行程。", bold_prefix="  - user_id (Integer, FK): ", indent=0.4)
    add_paragraph("行程標題（如 'Bali Bespoke Escape'）。", bold_prefix="  - trip_title (String): ", indent=0.4)
    add_paragraph("以長文字 (Text) 形式完整儲存 AI 生成之行程 JSON 資料，以利前端直接載入。", bold_prefix="  - itinerary_data (Text): ", indent=0.4)
    add_paragraph("建立時間戳記，便於列表按時間排序。", bold_prefix="  - created_at (DateTime): ", indent=0.4)

    # ---------------------------------------------------------------------------
    # 四、 核心功能實現與優化細節
    # ---------------------------------------------------------------------------
    add_heading_1("三、 核心功能實現與優化細節 (Features & Optimizations)")
    
    add_heading_2("1. 佈局重構與 Sticky 定位修復")
    add_paragraph(
        "我們調整了登入後的雙欄主工作台佈局。在 LG（大螢幕）下，左側行程 Timeline 佔用 65% 的寬度，"
        "提供寬敞的呼吸空間以網格呈現 Days 卡片；右側地圖則佔用 35% 寬度並設置為 lg:sticky 黏性固定。"
    )
    add_paragraph(
        "原先版面配置中，地圖無法隨畫面滾動跟隨。經過程式碼追蹤，我們發現登入後最外層的根容器 div 設置了 overflow-y-auto，"
        "導致其內部產生不可見的滾動阻尼，使 CSS Sticky 機制無法對焦 window 視窗。我們移除了該 overflow-y-auto 屬性，"
        "將全頁滾動交給 window 處理，成功使地圖容器（lg:sticky lg:top-6 lg:h-[calc(100vh-120px)]）在滾動時穩定懸浮於螢幕右側，"
        "使用者在閱讀長篇行程並點選景點時，地圖會即時以 flyTo 對焦，體驗一氣呵成。"
    )

    add_heading_2("2. 雙排版軸向切換功能 (Columns / Rows)")
    add_paragraph(
        "為了滿足不同的行程檢視偏好，本專案新增了「X/Y 軸排版對調」的切換功能，提供兩大檢視模式："
    )
    add_paragraph(
        "Days 卡片在寬螢幕下橫向並列為三欄，景點活動在各天卡片內以單欄垂直排列。適合需要一目了然看清「第幾天做什麼」的時間軸檢視。",
        bold_prefix="• 天數並排模式 (Columns, 預設): ", indent=0.3
    )
    add_paragraph(
        "天數卡片改為全寬一行一行向下垂直堆疊，而各天內部的景點則改為橫向網格排列（大螢幕為三欄並列）。這有效實現了 X 軸與 Y 軸的對換，"
        "讓單天的行程呈現出如同「景點瀑布流」般精緻的卡片橫向分佈，特別適合景點數量多、想要深入了解單天行程細節的旅客。",
        bold_prefix="• 天數堆疊模式 (Rows): ", indent=0.3
    )
    add_paragraph(
        "切換功能以一個高質感的毛玻璃 Switch 開關置於 Timeline checklist 標題旁，透過 React 狀態 `layoutAxis` 與 "
        "Tailwind 動態類別綁定，點擊時版面切換毫無延遲，且對既有的圖片懶加載與打卡狀態無任何副作用。"
    )

    add_heading_2("3. React 代碼重構與效能優化 (Performance Optimization)")
    add_paragraph("為了提升整個應用的反應速度並避免渲染卡頓，我們對 App.jsx 進行了全面的程式碼審查與重構：")
    add_paragraph(
        "將 `fetchSavedTrips` 的定義提升至一鍵儲存函數 `handleSaveItinerary` 之前，符合 JavaScript 嚴格的先宣告後使用（Hoisting）最佳實踐，"
        "並將 `fetchSavedTrips` 補齊至 `handleSaveItinerary` 的 `useCallback` 依賴陣列中，防止潛在的 Stale Closures（閉包過期）參考錯誤。",
        bold_prefix="• 健全依賴鏈: ", indent=0.3
    )
    add_paragraph(
        "移除了原先寫在登入頁面與導航列 Navbar 內部 button 上的 `onClick={() => setLang(...)}` 等 Inline 匿名箭頭函數，"
        "改以 React 快取的 `toggleLanguage` 與 `toggleLayoutAxis` 處理常式進行綁定。這能確保在每次 App 元件更新時，"
        "按鈕的 `onClick` 事件不會重新分配新的記憶體位址，進而避免子元件無謂的重複重繪。",
        bold_prefix="• 消除匿名函數參照: ", indent=0.3
    )
    add_paragraph(
        "使用 `useMemo` 快取了多語系字典字典查詢 `t`，僅在 `lang` 改變時才重新計算；而點擊地標、打卡景點等繁重邏輯均以 `useCallback` 封裝，"
        "使多卡片的大型渲染树運作更為高效。",
        bold_prefix="• 快取翻譯與 Callback: ", indent=0.3
    )

    # ---------------------------------------------------------------------------
    # 五、 專案總結與展望
    # ---------------------------------------------------------------------------
    add_heading_1("四、 專案總結與展望 (Summary & Outlook)")
    add_paragraph(
        "本階段開發成功克服了 LOKA 旅行規劃助理在排版交互與地圖定位上的重大痛點。透過巧妙的 CSS overflow 修正、"
        "創新的雙軸向切換設計，以及深度的 React Hooks 快取最佳實踐，LOKA 不僅在視覺美感上達到了奢華與輕量化並存的品牌調性，"
        "在運行效能與邏輯健壯性上也均達到了簡報發表與實際產品體驗的優異標準。"
    )
    add_paragraph(
        "展望未來，系統可進一步考慮串接真實的支付金流 API（如 Stripe），將 OSRM 導航拓展至多車道即時路況避堵，"
        "並引進 PDF 伺服器端渲染 (SSR) 技術，以提供更高像素的實體紙本行程列印。LOKA 將持續扮演奢華旅客探索世界的最佳貼身智慧助手。"
    )

    # ---------------------------------------------------------------------------
    # 存檔
    # ---------------------------------------------------------------------------
    output_path = "c:\\Users\\richa\\OneDrive\\Documents\\f3inalproject\\LOKA_Final_Project_Report.docx"
    doc.save(output_path)
    print(f"Report generated successfully at: {output_path}")

if __name__ == "__main__":
    create_report()
